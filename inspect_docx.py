import sys
from docx import Document

def inspect_docx(filepath):
    doc = Document(filepath)
    print("--- Paragraphs ---")
    for i, p in enumerate(doc.paragraphs[:30]):
        text = p.text.strip()
        if text:
            runs = p.runs
            run_info = []
            for r in runs:
                if not r.text.strip(): continue
                font = r.font.name or (p.style.font.name if p.style and p.style.font else None)
                size = r.font.size.pt if (r.font and r.font.size) else (p.style.font.size.pt if p.style and p.style.font and p.style.font.size else None)
                bold = r.bold or (p.style.font.bold if p.style and p.style.font else None)
                color = r.font.color.rgb if r.font and r.font.color and r.font.color.type == 1 else "Default"
                run_info.append(f"[{font}, {size}pt, b:{bold}, c:{color}]")
            
            style_name = p.style.name if p.style else "None"
            print(f"P{i} [{style_name}]: {text[:50]}... -> {', '.join(run_info)}")

    print("\n--- Tables ---")
    for t in doc.tables:
        print(f"Table (rows: {len(t.rows)}), Style: {t.style.name if t.style else 'None'}")
        if len(t.rows) > 0:
            header_cells = [c.text.strip().replace('\n', ' ') for c in t.rows[0].cells]
            print(f"  Row 0: {header_cells}")

inspect_docx("猫老师_历史记录测评报告 (15).docx")
