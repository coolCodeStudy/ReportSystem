import docx

doc = docx.Document("猫老师_历史记录测评报告.docx")

def safe_print(obj):
    try:
        print(obj)
    except:
        pass

for i, table in enumerate(doc.tables[:3]):
    print(f"\n--- Table {i} ---")
    
    for row_idx, row in enumerate(table.rows[:2]):
        print(f"Row {row_idx}:")
        for col_idx, cell in enumerate(row.cells):
            txt = cell.text.replace('\n', ' ')
            print(f"  Col {col_idx}: '{txt[:30]}'...")
            for p in cell.paragraphs:
                pf = p.paragraph_format
                print(f"    Para spacing_before={pf.space_before}, after={pf.space_after}, line_spacing={pf.line_spacing}, rule={pf.line_spacing_rule}")
                for r in p.runs:
                    print(f"      Run: text='{r.text[:10]}', size={r.font.size.pt if r.font and r.font.size else None}")
